import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls

def create_interview_guide_docx(output_path):
    doc = Document()

    # Set page margins (0.75 in)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    # Color Palette
    PRIMARY_NAVY = RGBColor(31, 78, 121)     # #1F4E79
    SECONDARY_BLUE = RGBColor(41, 128, 185)  # #2980B9
    DARK_TEXT = RGBColor(44, 62, 80)         # #2C3E50
    GRAY_TEXT = RGBColor(108, 117, 125)      # #6C757D
    BG_LIGHT_HEX = "F8F9FA"
    BORDER_BLUE_HEX = "1F4E79"

    # Helper function for setting shaded box background & left border
    def set_cell_background_and_border(cell, fill_hex, border_hex=None):
        tcPr = cell._tc.get_or_add_tcPr()
        # Shading
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
        tcPr.append(shd)
        
        if border_hex:
            # Set left border thick, clear top/right/bottom
            borders = parse_xml(f'''
                <w:tcBorders {nsdecls("w")}>
                    <w:left w:val="single" w:sz="24" w:space="0" w:color="{border_hex}"/>
                    <w:top w:val="none"/>
                    <w:right w:val="none"/>
                    <w:bottom w:val="none"/>
                </w:tcBorders>
            ''')
            tcPr.append(borders)

    # Title Banner
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("Retail Sales ETL & Analytics Platform\n")
    run_title.font.name = "Arial"
    run_title.font.size = Pt(24)
    run_title.font.bold = True
    run_title.font.color.rgb = PRIMARY_NAVY

    run_sub = p_title.add_run("Complete Interview Preparation & Q&A Guide\n")
    run_sub.font.name = "Arial"
    run_sub.font.size = Pt(14)
    run_sub.font.bold = True
    run_sub.font.color.rgb = SECONDARY_BLUE

    run_desc = p_title.add_run("Simple, beginner-friendly explanations & high-frequency MNC counter-questions (Accenture, Deloitte, TCS, Capgemini)")
    run_desc.font.name = "Arial"
    run_desc.font.size = Pt(10)
    run_desc.font.italic = True
    run_desc.font.color.rgb = GRAY_TEXT

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # Section Helper
    def add_heading_1(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.name = "Arial"
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.color.rgb = PRIMARY_NAVY
        return p

    def add_question(q_num, q_text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        run_num = p.add_run(f"Q{q_num}: ")
        run_num.font.name = "Arial"
        run_num.font.size = Pt(12)
        run_num.font.bold = True
        run_num.font.color.rgb = PRIMARY_NAVY

        run_txt = p.add_run(q_text)
        run_txt.font.name = "Arial"
        run_txt.font.size = Pt(12)
        run_txt.font.bold = True
        run_txt.font.color.rgb = DARK_TEXT
        return p

    def add_answer_box(answer_text):
        table = doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = table.cell(0, 0)
        cell.width = Inches(7.0)
        set_cell_background_and_border(cell, BG_LIGHT_HEX, BORDER_BLUE_HEX)
        
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.left_indent = Inches(0.15)
        p.paragraph_format.right_indent = Inches(0.15)
        
        run_lbl = p.add_run("ANSWER (How to explain in simple words):\n")
        run_lbl.font.name = "Arial"
        run_lbl.font.size = Pt(10)
        run_lbl.font.bold = True
        run_lbl.font.color.rgb = SECONDARY_BLUE
        
        run_ans = p.add_run(answer_text)
        run_ans.font.name = "Arial"
        run_ans.font.size = Pt(10.5)
        run_ans.font.color.rgb = DARK_TEXT
        
        doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # =========================================================================
    # SECTION 1
    # =========================================================================
    add_heading_1("Section 1: 60-Second Project Pitch & Business Value")

    add_question(1, "Tell me about a recent Data Engineering / Analytics project you have built.")
    add_answer_box(
        "\"I built an end-to-end Retail Sales ETL & Analytics Platform using Python, PostgreSQL, SQL, and Streamlit.\n\n"
        "The project takes raw, uncleaned retail sales data (around 10,000 transaction records), processes it through a 4-step automated Python ETL pipeline to validate and clean it, loads it into a PostgreSQL relational database, executes 20 SQL business queries, and visualizes the insights on a 4-page interactive Streamlit dashboard.\n\n"
        "I designed the pipeline with a modular architecture (Extract, Validate, Transform, Load), implemented error handling and logging, secured credentials using environment variables (.env), and added a smart fallback so the dashboard works even if PostgreSQL is offline.\""
    )

    add_question(2, "What business problem does this project solve?")
    add_answer_box(
        "\"Retail companies collect thousands of sales transactions across different cities, categories, and customer types. Raw data often has duplicates, missing values, and formatting errors—making it risky for management to trust.\n\n"
        "My project solves three main business problems:\n"
        "1. Data Quality: It automatically removes duplicate orders and cleans bad data.\n"
        "2. Profitability Analysis: It creates metrics like 'revenue' and 'profit_margin' so managers can instantly see which products make money and which ones cause losses.\n"
        "3. Executive Visibility: It provides an interactive dashboard where business managers can filter sales by region, category, or customer segment without writing any code or SQL queries.\""
    )

    # =========================================================================
    # SECTION 2
    # =========================================================================
    add_heading_1("Section 2: Architecture & Data Lifecycle")

    add_question(3, "Can you explain the technical architecture of your project?")
    add_answer_box(
        "\"Yes! The project is structured into 5 simple layers:\n\n"
        "1. Data Source Layer: Raw transaction dataset in CSV format (data/raw/retail_sales.csv).\n"
        "2. Python ETL Layer (scripts/):\n"
        "   - extract.py: Ingests raw CSV into Pandas.\n"
        "   - validate.py: Checks missing values, duplicates, and invalid numbers; saves a log report.\n"
        "   - transform.py: Deduplicates rows, fills null values, converts strings to Title Case, renames headers to snake_case, and calculates revenue & profit margin.\n"
        "   - load.py: Connects via SQLAlchemy and bulk-loads cleaned data into PostgreSQL.\n"
        "   - run_pipeline.py: Runs all 4 scripts in order with timing and persistent logging in logs/pipeline.log.\n"
        "3. Database Layer (database/):\n"
        "   - schema.sql: Defines the PostgreSQL table with proper data types.\n"
        "   - queries.sql: Contains 20 SQL queries for business KPIs.\n"
        "4. Presentation Layer (dashboard/):\n"
        "   - Interactive 4-page Streamlit web app built with Plotly charts.\n"
        "5. Deployment Layer:\n"
        "   - Hosted on GitHub and deployed on Streamlit Community Cloud with automatic CSV fallback.\""
    )

    add_question(4, "Walk me through the lifecycle of a single record from raw CSV to the final dashboard.")
    add_answer_box(
        "\"Let's trace one transaction row:\n\n"
        "1. Extract: extract.py reads a row: Ship Mode='Second Class', Sales='261.96', Profit='41.9136'.\n"
        "2. Validate: validate.py confirms Sales > 0 and text fields are non-empty.\n"
        "3. Transform: transform.py trims whitespace, changes header 'Ship Mode' to 'ship_mode', and calculates profit_margin = (41.9136 / 261.96) * 100 = 16.0%.\n"
        "4. Load: load.py inserts the row into PostgreSQL table 'sales' via SQLAlchemy.\n"
        "5. Visualize: dashboard/database.py queries the row and plots it on Streamlit charts!\""
    )

    # =========================================================================
    # SECTION 3
    # =========================================================================
    add_heading_1("Section 3: Detailed Module-by-Module Technical Deep Dive")

    add_question(5, "How did you implement the Extract module (scripts/extract.py)?")
    add_answer_box(
        "\"The Extract module safely ingests raw data:\n"
        "• check_file_exists(): Uses os.path.isfile() to verify data/raw/retail_sales.csv exists. Raises FileNotFoundError if missing.\n"
        "• load_dataset(): Uses pd.read_csv() inside try-except blocks to catch ParserError.\n"
        "• display_dataset_summary(): Prints total rows, columns, data types, and missing value counts.\n"
        "• extract_data(): Main function that orchestrates extraction and returns a Pandas DataFrame.\""
    )

    add_question(6, "What data quality checks did you implement in scripts/validate.py?")
    add_answer_box(
        "\"I implemented 6 automated checks:\n"
        "1. Missing Values: Counts NaN cells across all columns.\n"
        "2. Duplicates: Counts fully identical rows using df.duplicated().\n"
        "3. Sales Check: Ensures Sales >= 0.\n"
        "4. Quantity Check: Ensures Quantity >= 1 (order quantity cannot be zero or negative).\n"
        "5. Discount Check: Ensures Discount >= 0.\n"
        "6. Text Check: Ensures category and segment names are not empty.\n\n"
        "All results are printed to the terminal and saved to logs/validation_report.txt.\""
    )

    add_question(7, "What cleaning and feature engineering steps are done in scripts/transform.py?")
    add_answer_box(
        "\"Transformation is executed in 6 clean steps:\n"
        "1. Deduplication: Removes 17 duplicate rows using df.drop_duplicates() (rows reduced from 9,994 to 9,977).\n"
        "2. Missing Values: Fills missing numbers with 0.0 and missing text with 'Unknown'.\n"
        "3. Date Coercion: Converts date strings to datetime objects safely.\n"
        "4. Text Standardizing: Strips whitespace and converts text to Title Case.\n"
        "5. Snake_case Renaming: Renames column headers like 'Postal Code' to 'postal_code'.\n"
        "6. Feature Engineering:\n"
        "   - Creates 'revenue' (alias for sales).\n"
        "   - Creates 'profit_margin = np.where(sales != 0, ((profit / sales) * 100).round(2), 0.0)' to safely handle zero division.\""
    )

    add_question(8, "How does your Load module (scripts/load.py) interact with PostgreSQL?")
    add_answer_box(
        "\"The Load module performs secure database ingestion:\n"
        "• load_environment(): Reads credentials (host, port, db, user, password) from .env using python-dotenv.\n"
        "• create_connection(): Creates a SQLAlchemy engine with postgresql+psycopg2 and tests connection with SELECT 1.\n"
        "• load_dataframe(): Uses df.to_sql(name='sales', con=engine, if_exists='replace', index=False, method='multi') for fast batch insertion.\n"
        "• verify_load(): Runs SELECT COUNT(*) FROM sales to confirm all 9,977 rows were inserted successfully.\""
    )

    add_question(9, "How does scripts/run_pipeline.py orchestrate the ETL flow and handle errors?")
    add_answer_box(
        "\"run_pipeline.py is the master script:\n"
        "• It executes: run_extract() -> run_validation() -> run_transformation() -> run_load().\n"
        "• Error Handling: If any step fails, it catches the error, logs the failure, prints a clear error message, and stops execution immediately using sys.exit(1).\n"
        "• Logging: Uses Python's built-in logging module to append timestamped logs (INFO/ERROR, step duration, status) to logs/pipeline.log.\""
    )

    add_question(10, "How is the Streamlit Dashboard organized and how does caching work?")
    add_answer_box(
        "\"The dashboard is organized into 5 clean files:\n"
        "• app.py: Main entry point rendering 4 pages (Executive, Sales, Product, Customer).\n"
        "• database.py: Handles database loading with a smart fallback mechanism.\n"
        "• charts.py: Contains pure Plotly chart functions.\n"
        "• components.py: Renders sidebar filters and 5 headline KPI metric cards.\n"
        "• utils.py: Formatting ($1.2M, 12.0%) and filtering helpers.\n\n"
        "Caching Strategy:\n"
        "• @st.cache_resource: Keeps the database engine connection alive across user sessions.\n"
        "• @st.cache_data(ttl=300): Caches query results for 5 minutes. Clicking 'Refresh Data' clears the cache to pull fresh data.\""
    )

    # =========================================================================
    # SECTION 4
    # =========================================================================
    add_heading_1("Section 4: High-Frequency Counter-Questions (Accenture / MNC Style)")

    add_question(11, "Why did you use SQLAlchemy instead of standard psycopg2 cursors?")
    add_answer_box(
        "\"SQLAlchemy offers two big advantages:\n"
        "1. Direct Pandas Integration: Pandas to_sql() works seamlessly with SQLAlchemy engines, allowing fast batch insertion without writing manual SQL INSERT loops.\n"
        "2. Connection Pooling: SQLAlchemy manages connection pooling and automatic reconnections (pool_pre_ping=True), keeping web applications like Streamlit stable.\""
    )

    add_question(12, "What is the difference between if_exists='replace' and if_exists='append' in to_sql()?")
    add_answer_box(
        "\"In Pandas to_sql():\n"
        "• if_exists='replace': Drops the existing table and recreates it before inserting records.\n"
        "• if_exists='append': Keeps the existing table and inserts new rows at the end.\n\n"
        "In our batch project, I used 'replace' to ensure a clean state on every pipeline run. In a production pipeline with daily incremental loads, we would use 'append' along with staging tables.\""
    )

    add_question(13, "Why did you use np.where() when calculating profit margin?")
    add_answer_box(
        "\"If a row has sales = 0, calculating (profit / sales) causes a ZeroDivisionError or outputs NaN/infinity.\n\n"
        "Using np.where(sales != 0, ((profit / sales) * 100).round(2), 0.0) checks the sales column first. If sales is non-zero, it calculates the margin; if sales is zero, it safely sets profit_margin to 0.0.\""
    )

    add_question(14, "Why did you create a fallback mechanism in dashboard/database.py?")
    add_answer_box(
        "\"When deploying dashboards to cloud platforms like Streamlit Community Cloud, a local PostgreSQL database won't be reachable.\n\n"
        "To make the dashboard 100% resilient:\n"
        "1. It tries connecting to PostgreSQL via .env.\n"
        "2. If database is offline, it falls back to data/processed/retail_sales_cleaned.csv.\n"
        "3. If processed CSV is missing, it runs the transformation script on data/raw/retail_sales.csv on the fly.\n"
        "This guarantees the dashboard always loads cleanly for users and recruiters.\""
    )

    add_question(15, "How would you scale this pipeline if data volume grew from 10,000 rows to 100 Million rows (100 GB)?")
    add_answer_box(
        "\"If data scales to 100 GB, in-memory Pandas processing will run out of memory (OOM). I would scale it as follows:\n"
        "1. Processing: Replace Pandas with PySpark or DuckDB for distributed parallel processing.\n"
        "2. Storage: Store raw files in cloud storage (AWS S3 or Azure Data Lake) using partitioned Parquet format.\n"
        "3. Database: Replace local PostgreSQL with a cloud data warehouse like Snowflake or AWS Redshift.\n"
        "4. Orchestration: Replace Python execution scripts with Apache Airflow DAGs for scheduling and retry monitoring.\""
    )

    add_question(16, "Why keep database credentials in .env instead of Python files?")
    add_answer_box(
        "\"Hardcoding passwords in code is a major security risk because code is pushed to version control like GitHub.\n\n"
        "By placing credentials in .env and adding .env to .gitignore, passwords stay safely on the server. In production, credentials are injected via secret managers like AWS Secrets Manager or Streamlit Secrets.\""
    )

    # =========================================================================
    # SECTION 5
    # =========================================================================
    add_heading_1("Section 5: Scenario & Troubleshooting Questions")

    add_question(17, "What happens if a raw CSV file contains duplicate rows? How does your code handle it?")
    add_answer_box(
        "\"In scripts/transform.py, remove_duplicates() calls df.drop_duplicates().reset_index(drop=True).\n\n"
        "In our dataset, it detected 17 duplicate rows. It kept the first record, removed the 17 identical copies, reset row index, and logged the result. Row count dropped cleanly from 9,994 to 9,977.\""
    )

    add_question(18, "What happens if scripts/run_pipeline.py finds a missing raw CSV file?")
    add_answer_box(
        "\"When step 1 (Extract) runs, check_file_exists() verifies the file path. If missing, it raises a FileNotFoundError.\n\n"
        "The try-except block in run_pipeline.py catches the error, logs 'STATUS: FAILED' in logs/pipeline.log, prints an error banner, and exits with sys.exit(1). Steps 2, 3, and 4 do not run.\""
    )

    add_question(19, "Name 3 key SQL queries from database/queries.sql and explain their business purpose.")
    add_answer_box(
        "1. Top 10 Sub-Categories by Revenue:\n"
        "   SELECT sub_category, SUM(revenue) AS total_revenue FROM sales GROUP BY sub_category ORDER BY total_revenue DESC LIMIT 10;\n"
        "   Business Purpose: Identifies top-selling products for inventory allocation.\n\n"
        "2. Bottom 5 Sub-Categories by Profit (Loss Makers):\n"
        "   SELECT sub_category, SUM(profit) AS total_profit FROM sales GROUP BY sub_category ORDER BY total_profit ASC LIMIT 5;\n"
        "   Business Purpose: Identifies items causing financial losses (e.g. Tables or Bookcases due to heavy discounts).\n\n"
        "3. Revenue & Profit Margin by Customer Segment:\n"
        "   SELECT segment, SUM(revenue) AS total_revenue, AVG(profit_margin) AS avg_margin FROM sales GROUP BY segment;\n"
        "   Business Purpose: Evaluates performance across Consumer, Corporate, and Home Office segments.\""
    )

    add_question(20, "What was the most challenging technical part of this project and how did you solve it?")
    add_answer_box(
        "\"The most challenging part was building a seamless data pipeline that works both locally with PostgreSQL and on cloud hosting (Streamlit Cloud) without breaking.\n\n"
        "Cloud environments don't have access to local PostgreSQL databases. I solved this by engineering a 3-tier data access strategy in dashboard/database.py. It automatically checks for PostgreSQL, falls back to the cleaned CSV, or generates data on the fly. This guaranteed 100% dashboard uptime for recruiters while retaining full PostgreSQL capabilities locally.\""
    )

    # =========================================================================
    # REVISION CHEAT-SHEET TABLE
    # =========================================================================
    add_heading_1("Summary Cheat-Sheet for Interview Day")

    table_data = [
        ("Key Metric / Parameter", "Value / Detail"),
        ("Raw Dataset Size", "9,994 rows, 13 columns"),
        ("Cleaned Dataset Size", "9,977 rows (17 duplicates removed), 15 columns"),
        ("Engineered Features", "revenue (alias for sales), profit_margin (%)"),
        ("Total Revenue", "$2,297,200.86 ($2.30 Million)"),
        ("Total Profit", "$286,241.42 ($286.2 Thousand)"),
        ("Average Profit Margin", "12.01%"),
        ("ETL Modules", "extract.py, validate.py, transform.py, load.py, run_pipeline.py"),
        ("Dashboard Stack", "Streamlit, Plotly Express, Plotly Graph Objects"),
        ("SQL Queries Count", "20 Business KPI Queries (database/queries.sql)"),
        ("Dashboard Pages", "Executive Dashboard, Sales Performance, Product Analysis, Customer Insights"),
        ("Live Streamlit URL", "https://retail-sales-data-platform.streamlit.app/"),
        ("GitHub Repository", "https://github.com/mahimayadav04/retail-sales-etl-analytics"),
    ]

    t = doc.add_table(rows=len(table_data), cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    for row_idx, (col1_txt, col2_txt) in enumerate(table_data):
        row_cells = t.rows[row_idx].cells
        row_cells[0].width = Inches(2.5)
        row_cells[1].width = Inches(4.5)
        
        p0 = row_cells[0].paragraphs[0]
        p1 = row_cells[1].paragraphs[0]
        p0.paragraph_format.space_before = Pt(4)
        p0.paragraph_format.space_after = Pt(4)
        p1.paragraph_format.space_before = Pt(4)
        p1.paragraph_format.space_after = Pt(4)

        if row_idx == 0:
            set_cell_background_and_border(row_cells[0], "1F4E79")
            set_cell_background_and_border(row_cells[1], "1F4E79")
            r0 = p0.add_run(col1_txt)
            r0.font.bold = True
            r0.font.color.rgb = RGBColor(255, 255, 255)
            r1 = p1.add_run(col2_txt)
            r1.font.bold = True
            r1.font.color.rgb = RGBColor(255, 255, 255)
        else:
            bg_color = "F8F9FA" if row_idx % 2 == 1 else "FFFFFF"
            set_cell_background_and_border(row_cells[0], bg_color)
            set_cell_background_and_border(row_cells[1], bg_color)
            
            r0 = p0.add_run(col1_txt)
            r0.font.bold = True
            r0.font.color.rgb = DARK_TEXT
            
            r1 = p1.add_run(col2_txt)
            r1.font.color.rgb = DARK_TEXT

    doc.save(output_path)
    print(f"[OK] Created Word document at: {output_path}")

if __name__ == "__main__":
    out_file = os.path.join("d:", os.sep, "Retail-Sales-ETL-Analytics", "Retail_Sales_ETL_Interview_Guide.docx")
    create_interview_guide_docx(out_file)
